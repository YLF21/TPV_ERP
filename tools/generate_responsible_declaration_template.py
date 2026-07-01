from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "documentos" / (
    "Plantilla_Declaracion_Responsable_TPV_ERP.docx"
)

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
FIELD_FILL = "FFF2CC"
LIGHT_GRAY = "F2F2F2"
TEXT_GRAY = RGBColor(89, 89, 89)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_run(run, color=FIELD_FILL):
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    r_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header(row):
    set_repeat_table_header(row)


def set_font(run, name="Aptos", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_field(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    set_font(run, bold=bold)
    shade_run(run)
    return run


def style_paragraph(paragraph, before=0, after=6, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_label_field(document, label, placeholder):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.1)
    table.columns[1].width = Cm(11.2)
    label_cell, value_cell = table.rows[0].cells
    label_cell.width = Cm(5.1)
    value_cell.width = Cm(11.2)
    set_cell_fill(label_cell, LIGHT_BLUE)
    set_cell_fill(value_cell, FIELD_FILL)
    for cell in table.rows[0].cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    label_p = label_cell.paragraphs[0]
    value_p = value_cell.paragraphs[0]
    style_paragraph(label_p, after=0)
    style_paragraph(value_p, after=0)
    set_font(label_p.add_run(label), bold=True, color=RGBColor(31, 78, 120))
    set_font(value_p.add_run(placeholder))
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, before=8, after=4)
    run = paragraph.add_run(text)
    set_font(run, size=11.5, bold=True, color=RGBColor(31, 78, 120))
    return paragraph


def add_long_field(document, label, placeholder, min_lines=3):
    add_section_heading(document, label)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_fill(cell, FIELD_FILL)
    set_cell_margins(cell, top=140, bottom=140)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line=1.15)
    set_font(paragraph.add_run(placeholder))
    for _ in range(min_lines - 1):
        paragraph.add_run("\n")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_yes_no(document, label, guidance=None):
    add_section_heading(document, label)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, text in enumerate(("☐ S - Sí", "☐ N - No")):
        cell = table.cell(0, index)
        cell.width = Cm(8.15)
        set_cell_fill(cell, FIELD_FILL)
        set_cell_margins(cell, top=130, bottom=130)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0)
        set_font(p.add_run(text), bold=True)
    if guidance:
        p = document.add_paragraph()
        style_paragraph(p, after=5)
        set_font(p.add_run(guidance), size=9, color=TEXT_GRAY)


def add_footer(section):
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Cm(16.3))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(11.5)
    right.width = Cm(4.8)
    left_p = left.paragraphs[0]
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(left_p.add_run("TPV ERP | Declaración responsable | Versión: [VERSIÓN]"), size=8, color=TEXT_GRAY)
    set_font(right_p.add_run("Página "), size=8, color=TEXT_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = right_p.add_run()
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, after=4)
    set_font(
        title.add_run("DECLARACIÓN RESPONSABLE DEL SISTEMA\nINFORMÁTICO DE FACTURACIÓN"),
        size=16,
        bold=True,
        color=RGBColor(31, 78, 120),
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, after=12)
    set_font(subtitle.add_run("Plantilla editable para TPV ERP"), size=10, bold=True, color=TEXT_GRAY)

    notice = document.add_table(rows=1, cols=1)
    notice.alignment = WD_TABLE_ALIGNMENT.CENTER
    notice_cell = notice.cell(0, 0)
    set_cell_fill(notice_cell, LIGHT_GRAY)
    set_cell_margins(notice_cell, top=140, bottom=140, start=160, end=160)
    p = notice_cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.12)
    set_font(p.add_run("NOTA DE USO. "), bold=True, color=RGBColor(31, 78, 120))
    set_font(
        p.add_run(
            "Complete todos los campos entre corchetes para una versión concreta. "
            "La modalidad y las características declaradas deben coincidir con el "
            "producto realmente distribuido. Debe conservarse una declaración por "
            "cada versión. Esta plantilla no acredita por sí sola el cumplimiento."
        ),
        size=9.5,
    )

    document.add_paragraph()
    add_label_field(
        document,
        "1.a) Nombre del sistema informático",
        "[TPV ERP / DENOMINACIÓN COMERCIAL DEFINITIVA]",
    )
    add_label_field(
        document,
        "1.b) Código identificador del sistema",
        "[CÓDIGO BREVE Y UNÍVOCO, POR EJEMPLO: TPVERP]",
    )
    add_label_field(
        document,
        "1.c) Identificador completo de versión",
        "[VERSIÓN COMPLETA DECLARADA, POR EJEMPLO: 1.0.0+2026.06]",
    )
    add_long_field(
        document,
        "1.d) Componentes hardware y software, descripción y funcionalidades",
        "[DESCRIBIR LOS COMPONENTES DE LA INSTALACIÓN: servidor o equipo local, "
        "terminales, sistema operativo, aplicación TPV ERP, base de datos, módulos "
        "de facturación y componentes auxiliares. Resumir las principales funciones "
        "de facturación, emisión de facturas completas y simplificadas, conservación "
        "de registros, trazabilidad, exportación y comunicaciones aplicables.]",
        min_lines=4,
    )
    add_yes_no(
        document,
        "1.e) ¿El sistema solo puede funcionar exclusivamente como VERI*FACTU?",
        "Marque una sola opción conforme al comportamiento real de esta versión.",
    )
    add_yes_no(
        document,
        "1.f) ¿Permite su uso por varios obligados tributarios?",
        "Incluye el uso por un mismo usuario para la facturación de varios obligados tributarios.",
    )
    add_long_field(
        document,
        "1.g) Tipos de firma utilizados en registros de facturación y de evento",
        "[SI LA VERSIÓN NO ES EXCLUSIVAMENTE VERI*FACTU, IDENTIFICAR LOS TIPOS Y "
        "ALGORITMOS DE FIRMA UTILIZADOS. SI ES EXCLUSIVAMENTE VERI*FACTU, INDICAR "
        "QUE NO PROCEDE Y EXPLICAR BREVEMENTE EL MOTIVO.]",
        min_lines=3,
    )
    add_label_field(
        document,
        "1.h) Productor: nombre, apellidos o razón social",
        "[NOMBRE COMPLETO O RAZÓN SOCIAL DEL PRODUCTOR]",
    )
    add_label_field(
        document,
        "1.i) NIF o identificación equivalente",
        "[NIF ESPAÑOL / TIPO, NÚMERO Y PAÍS DE LA IDENTIFICACIÓN]",
    )
    add_long_field(
        document,
        "1.j) Dirección postal completa de contacto",
        "[CALLE, NÚMERO, PISO, CÓDIGO POSTAL, LOCALIDAD, PROVINCIA Y PAÍS]",
        min_lines=2,
    )

    add_section_heading(document, "1.k) Manifestación de cumplimiento normativo")
    compliance = document.add_table(rows=1, cols=1)
    compliance.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = compliance.cell(0, 0)
    set_cell_fill(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=160, bottom=160, start=160, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_paragraph(p, after=0, line=1.15)
    set_font(
        p.add_run(
            "La persona o entidad productora identificada en esta declaración hace "
            "constar que el sistema informático indicado, en la versión señalada, "
            "cumple con lo dispuesto en el artículo 29.2.j) de la Ley 58/2003, de "
            "17 de diciembre, General Tributaria; en el Reglamento aprobado por el "
            "Real Decreto 1007/2023, de 5 de diciembre; en la Orden HAC/1177/2024, "
            "de 17 de octubre; y en las especificaciones publicadas en la sede "
            "electrónica de la Agencia Estatal de Administración Tributaria que "
            "completen dicha orden."
        )
    )
    p = document.add_paragraph()
    style_paragraph(p, after=8)
    add_field(p, "☐ Confirmo esta manifestación para la versión indicada.", bold=True)

    add_section_heading(document, "1.l) Fecha y lugar de suscripción")
    signing = document.add_table(rows=2, cols=2)
    signing.alignment = WD_TABLE_ALIGNMENT.CENTER
    signing.autofit = False
    values = (
        ("Fecha completa", "[DÍA] de [MES] de [AÑO]"),
        ("Lugar", "[LOCALIDAD, PROVINCIA] - [PAÍS]"),
    )
    for row, (label, value) in zip(signing.rows, values):
        row.cells[0].width = Cm(5.1)
        row.cells[1].width = Cm(11.2)
        set_cell_fill(row.cells[0], LIGHT_BLUE)
        set_cell_fill(row.cells[1], FIELD_FILL)
        for current in row.cells:
            set_cell_margins(current)
            current.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=RGBColor(31, 78, 120))
        set_font(row.cells[1].paragraphs[0].add_run(value))

    document.add_paragraph()
    signature = document.add_table(rows=4, cols=2)
    signature.alignment = WD_TABLE_ALIGNMENT.CENTER
    signature.autofit = False
    signature_data = (
        ("Nombre de quien firma", "[NOMBRE Y APELLIDOS]"),
        ("Cargo o representación", "[CARGO / TÍTULO DE REPRESENTACIÓN]"),
        ("Firma", "\n\n[FIRMA MANUSCRITA O ELECTRÓNICA]\n"),
        ("Sello, si procede", "\n[SELLO]\n"),
    )
    for row, (label, value) in zip(signature.rows, signature_data):
        row.cells[0].width = Cm(5.1)
        row.cells[1].width = Cm(11.2)
        set_cell_fill(row.cells[0], LIGHT_GRAY)
        set_cell_fill(row.cells[1], FIELD_FILL)
        for current in row.cells:
            set_cell_margins(current, top=120, bottom=120)
            current.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(row.cells[0].paragraphs[0].add_run(label), bold=True)
        set_font(row.cells[1].paragraphs[0].add_run(value))

    document.add_section(WD_SECTION.NEW_PAGE)
    add_footer(document.sections[0])
    add_footer(document.sections[1])

    annex_title = document.add_paragraph()
    annex_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(annex_title, after=10)
    set_font(annex_title.add_run("ANEXO RECOMENDADO"), size=15, bold=True, color=RGBColor(31, 78, 120))
    p = document.add_paragraph()
    style_paragraph(p, after=10)
    set_font(
        p.add_run(
            "Información adicional del productor, del sistema y del acceso al histórico "
            "de declaraciones responsables."
        )
    )
    add_label_field(document, "2.a) Teléfono de contacto", "[TELÉFONO]")
    add_label_field(document, "2.a) Correo electrónico", "[CORREO ELECTRÓNICO]")
    add_label_field(document, "2.b) Sitio web del productor", "[URL]")
    add_label_field(document, "2.b) Información del producto", "[URL DE TPV ERP]")
    add_label_field(
        document,
        "2.b) Histórico de declaraciones",
        "[URL O UBICACIÓN DONDE SE CONSERVAN Y CONSULTAN LAS DECLARACIONES]",
    )
    add_long_field(
        document,
        "2.c) Información técnica adicional",
        "[DESCRIBIR ARQUITECTURA, TIPO DE INSTALACIÓN, DEPENDENCIAS, MODALIDAD DE "
        "FUNCIONAMIENTO, MECANISMOS DE REGISTRO, ENCADENAMIENTO, HUELLA, FIRMA, "
        "CONSERVACIÓN, EXPORTACIÓN Y COMUNICACIÓN QUE RESULTEN APLICABLES.]",
        min_lines=6,
    )
    add_long_field(
        document,
        "Control interno de la versión declarada",
        "[FECHA DE PUBLICACIÓN, IDENTIFICADOR DE COMPILACIÓN, REPOSITORIO O ARTEFACTO "
        "VERIFICADO, RESPONSABLE DE LA REVISIÓN Y EVIDENCIAS DE PRUEBA CONSERVADAS.]",
        min_lines=4,
    )

    refs = document.add_paragraph()
    style_paragraph(refs, before=12, after=4)
    set_font(refs.add_run("Referencias normativas"), size=11.5, bold=True, color=RGBColor(31, 78, 120))
    for text in (
        "Real Decreto 1007/2023, de 5 de diciembre, artículo 13.",
        "Orden HAC/1177/2024, de 17 de octubre, artículo 15.",
        "Especificaciones y ejemplos publicados por la Agencia Estatal de Administración Tributaria.",
    ):
        p = document.add_paragraph(style="List Bullet")
        style_paragraph(p, after=3)
        set_font(p.add_run(text), size=9.5)

    document.core_properties.title = "Plantilla de declaración responsable del sistema informático de facturación"
    document.core_properties.subject = "TPV ERP - Declaración responsable del fabricante"
    document.core_properties.author = "TPV ERP"
    document.core_properties.keywords = "TPV ERP, VERI*FACTU, declaración responsable, fabricante"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
